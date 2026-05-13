# -*- coding: utf-8 -*-
from __future__ import annotations

import math
import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from docx import Document


ROOT = Path(r"F:/data_set_process/data_process/05_rebuild_from_raw_20260511")
PROJECT_ROOT = Path(r"F:/data_set_process/data_process")
RAW_ROOT = PROJECT_ROOT / "01_datasets" / "数据预处理"
ROAD_CFG_DIR = PROJECT_ROOT / "01_datasets" / "多模态数据" / "被试数据集合" / "道路信息" / "道路"
SMALL_PAPER = Path(r"F:/开题报告/小论文_按修改意见修改稿_三线表_格式修正版.docx")

PREV_TABLE_DIR = ROOT / "02_samples" / "road_event_anchor_audit_v0_1" / "tables"
SESSION_MODULE_TIMES = PREV_TABLE_DIR / "session_module_entry_exit_v0_1.csv"
ROAD_POSITION_MAP = PREV_TABLE_DIR / "road_event_position_map_v0_1.csv"
SCENE_TRIGGER_SESSION_TIMES = (
    ROOT / "02_samples" / "scene_trigger_audit_v0_2" / "tables" / "scene_trigger_session_times_v0_2.csv"
)

OUT_DIR = ROOT / "02_samples" / "ego_direction_design_anchor_v0_4"
TABLE_DIR = OUT_DIR / "tables"
REPORT_DIR = ROOT / "09_reports"

CFG_MODULE_MAP = {
    "curve1": "curve1_Area2.cfg",
    "curve2": "curve2_Area2.cfg",
    "curve3": "curve3_Area2.cfg",
    "differentmu_road": "differentmu_road_Area2.cfg",
    "fix_road": "fix_road_Area2.cfg",
    "longstraight": "longsrtaight_Area2.cfg",
    "middle_section": "middle_section_Area2.cfg",
    "stop": "stop_Area2.cfg",
    "zd": "zd_Area2.cfg",
}

MODULE_RULES = {
    "middle_section": {
        "design_category_cn": "连续超车负荷事件段",
        "paper_evidence_cn": "middle_section 重复 9 次；交通车约 23 m/s；用户确认道路连接段存在连续超车事件",
        "primary_anchor_rule_cn": "连接段入口作为负荷开始；段内横向偏移变化、横摆角速度、横向加速度用于确认连续超车/变道响应",
        "anchor_status_cn": "主候选/连续负荷",
    },
    "longstraight": {
        "design_category_cn": "大货车紧急变道/直道交通交互",
        "paper_evidence_cn": "小论文：大货车紧急变道，目标车侵入本车道；.aed 中 MAN TGL truck 存在 ChangeLane 25->26，Chrysler300 存在 Stop；用户确认 longstraight 有变道触发点",
        "primary_anchor_rule_cn": "优先使用显式 ChangeLane/Stop 触发点作为场景设计候选，再用车身横向/制动响应确认；不能简单降级为背景车流",
        "anchor_status_cn": "主候选/显式触发",
    },
    "curve1": {
        "design_category_cn": "弯道路段",
        "paper_evidence_cn": "弯道路段 curve1、curve2、curve3；风险来源为道路曲率变化；锚点依据为局部 roll peak",
        "primary_anchor_rule_cn": "道路几何入口/曲率峰值作为候选，车身横滚或横向加速度峰值作为确认",
        "anchor_status_cn": "主候选",
    },
    "curve2": {
        "design_category_cn": "弯道路段",
        "paper_evidence_cn": "弯道路段 curve1、curve2、curve3；风险来源为道路曲率变化；锚点依据为局部 roll peak",
        "primary_anchor_rule_cn": "道路几何入口/曲率峰值作为候选，车身横滚或横向加速度峰值作为确认",
        "anchor_status_cn": "主候选",
    },
    "curve3": {
        "design_category_cn": "弯道路段",
        "paper_evidence_cn": "弯道路段 curve1、curve2、curve3；风险来源为道路曲率变化；锚点依据为局部 roll peak",
        "primary_anchor_rule_cn": "道路几何入口/曲率峰值作为候选，车身横滚或横向加速度峰值作为确认",
        "anchor_status_cn": "主候选但当前记录映射不足",
    },
    "differentmu_road": {
        "design_category_cn": "低附着路段",
        "paper_evidence_cn": "低附着路段；μ=0.8-0.2 不同附着区域；锚点依据为进入低 μ 区域",
        "primary_anchor_rule_cn": "以原始车辆 zx1|mu 下降点和 cfg 中低 μ 段入口/变化点作为候选",
        "anchor_status_cn": "主候选",
    },
    "fix_road": {
        "design_category_cn": "施工/维修变道/道路受限",
        "paper_evidence_cn": "小论文：施工/维修变道，车道收窄与交通干扰；.aed 中 MAN TGL 25->26 和 BMW m340 26->25 存在 ChangeLane 触发；用户确认维修路段涉及变道触发点",
        "primary_anchor_rule_cn": "优先使用 .aed 显式 ChangeLane 触发点和位置触发，再用车身姿态确认变道避让/速度调整",
        "anchor_status_cn": "主候选/显式触发",
    },
    "stop": {
        "design_category_cn": "前向急停",
        "paper_evidence_cn": "前车减速或停止；风险来源为 TTC 快速下降；锚点依据为显式触发/steer-rate peak80",
        "primary_anchor_rule_cn": "优先找被试方向前车急停触发；若缺失，则用制动/纵向减速度作为确认线索",
        "anchor_status_cn": "待确认",
    },
    "zd": {
        "design_category_cn": "匝道/汇入交互",
        "paper_evidence_cn": "匝道/汇入；多车道汇入与交通交互；锚点依据为汇入区/steer-rate peak80",
        "primary_anchor_rule_cn": "先确认 zd 是否对应匝道/汇入，再用汇入区位置和车身姿态确认",
        "anchor_status_cn": "待确认",
    },
}

VEHICLE_COLS = [
    "StorageTime",
    "zx1|v_km/h",
    "zx1|mu",
    "zx1|lateraldistance",
    "zx|BrakePedal",
    "zx|SteeringWheel",
    "zx|ax",
    "zx|ay",
    "zx|vyaw",
    "zx|vroll",
]


def ensure_dirs() -> None:
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_DIR.mkdir(parents=True, exist_ok=True)


def read_csv(path: Path, **kwargs: Any) -> pd.DataFrame:
    return pd.read_csv(path, encoding="utf-8-sig", low_memory=False, **kwargs)


def write_csv(df: pd.DataFrame, path: Path) -> None:
    df.to_csv(path, index=False, encoding="utf-8-sig")


def now_str() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def finite_float(value: Any, default: float = float("nan")) -> float:
    try:
        out = float(value)
    except (TypeError, ValueError):
        return default
    return out if math.isfinite(out) else default


def to_seconds(storage_time: pd.Series) -> np.ndarray:
    parsed = pd.to_datetime(storage_time, errors="coerce")
    out = np.full(len(storage_time), np.nan, dtype=np.float64)
    valid = parsed.notna().to_numpy()
    if valid.any():
        ns = parsed[valid].astype("datetime64[ns]").astype("int64").to_numpy(dtype=np.float64)
        out[valid] = ns / 1e9
    return out


def extract_small_paper_design() -> tuple[pd.DataFrame, str]:
    doc = Document(str(SMALL_PAPER))
    rows: list[dict[str, Any]] = []
    markdown: list[str] = [
        "# 小论文场景设计依据摘录 v0.4",
        "",
        f"来源文件：`{SMALL_PAPER}`",
        f"提取时间：{now_str()}",
        "",
        "## 正文关键段落",
        "",
    ]
    keywords = ["路线", "极端工况", "场景", "触发", "锚点", "低附着", "弯道", "急停", "汇入", "车道"]
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text and any(key in text for key in keywords):
            markdown.append(f"- P{idx}: {text}")
    markdown.append("")
    markdown.append("## 表格摘录")
    markdown.append("")
    for table_idx, table in enumerate(doc.tables):
        table_rows = [[cell.text.strip().replace("\n", " | ") for cell in row.cells] for row in table.rows]
        flat = "\n".join(" || ".join(row) for row in table_rows)
        if not any(key in flat for key in keywords):
            continue
        headers = table_rows[0] if table_rows else []
        markdown.append(f"### TABLE {table_idx}")
        markdown.append("")
        for row_idx, row in enumerate(table_rows):
            markdown.append("- " + " | ".join(row))
            if row_idx == 0:
                continue
            row_dict = {
                "source_docx": str(SMALL_PAPER),
                "table_index": table_idx,
                "row_index": row_idx,
            }
            for col_idx, value in enumerate(row):
                header = headers[col_idx] if col_idx < len(headers) and headers[col_idx] else f"col_{col_idx}"
                row_dict[header] = value
            rows.append(row_dict)
        markdown.append("")
    return pd.DataFrame(rows), "\n".join(markdown)


def parse_cfg_lane_geometry() -> tuple[pd.DataFrame, pd.DataFrame]:
    lane_rows: list[dict[str, Any]] = []
    for module_name, cfg_name in CFG_MODULE_MAP.items():
        path = ROAD_CFG_DIR / cfg_name
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8", errors="ignore")
        order = 0
        for match in re.finditer(r"(Straight|Circle)\s+l(\d+)\s*\{(?P<body>.*?)\n\s*\};", text, re.S):
            order += 1
            lane_type = match.group(1)
            lane_id = match.group(2)
            body = match.group("body")

            def grab(name: str) -> float:
                m = re.search(rf"{re.escape(name)}\s*=\s*([-\d.eE]+)", body)
                return finite_float(m.group(1)) if m else float("nan")

            length_match = re.search(r"#\s*Length\s*=\s*([-\d.eE]+)", body)
            length_m = finite_float(length_match.group(1)) if length_match else float("nan")
            direction_raw = str(int(grab("Direction"))) if math.isfinite(grab("Direction")) else ""
            dist_mid = grab("DistToMid0")
            if not math.isfinite(dist_mid):
                dist_mid = grab("DistToRef0")
            width = grab("Width0")
            lane_rows.append(
                {
                    "module_name": module_name,
                    "source_cfg": str(path),
                    "order_in_cfg": order,
                    "lane_type": lane_type,
                    "lane_id": lane_id,
                    "x0": grab("x0"),
                    "y0": grab("y0"),
                    "x1": grab("x1"),
                    "y1": grab("y1"),
                    "length_m": length_m,
                    "dist_to_mid_m": dist_mid,
                    "width_m": width,
                    "mu": grab("mu"),
                    "direction_raw": direction_raw,
                    "is_center_or_separator": abs(dist_mid) < 1e-6 if math.isfinite(dist_mid) else False,
                    "is_assumed_ego_direction": direction_raw == "1" and (abs(dist_mid) > 1e-6 if math.isfinite(dist_mid) else True),
                }
            )
    lanes = pd.DataFrame(lane_rows)
    if lanes.empty:
        return lanes, lanes.copy()

    drive = lanes[(lanes["is_assumed_ego_direction"]) & (lanes["width_m"].fillna(0) >= 3.0)].copy()
    if drive.empty:
        return lanes, pd.DataFrame()

    group_cols = ["module_name", "direction_raw", "x0", "y0", "x1", "y1", "length_m"]
    segment_rows: list[dict[str, Any]] = []
    for key, group in drive.groupby(group_cols, dropna=False, sort=False):
        module = str(key[0])
        segment_rows.append(
            {
                "module_name": module,
                "direction_raw": key[1],
                "segment_order_in_cfg": int(group["order_in_cfg"].min()),
                "lane_ids": ";".join(group["lane_id"].astype(str).tolist()),
                "mu_values": ";".join(str(v) for v in sorted(group["mu"].dropna().unique())),
                "min_mu": float(group["mu"].min()),
                "max_mu": float(group["mu"].max()),
                "segment_length_m": float(group["length_m"].median()),
                "x0": key[2],
                "y0": key[3],
                "x1": key[4],
                "y1": key[5],
            }
        )
    segments = pd.DataFrame(segment_rows).sort_values(["module_name", "segment_order_in_cfg"]).reset_index(drop=True)
    cum_rows: list[dict[str, Any]] = []
    for module, group in segments.groupby("module_name", sort=False):
        start = 0.0
        for _, row in group.iterrows():
            length = finite_float(row["segment_length_m"], 0.0)
            out = row.to_dict()
            out["module_relative_s_start_m"] = round(start, 3)
            out["module_relative_s_end_m"] = round(start + length, 3)
            out["module_relative_mid_m"] = round(start + length / 2.0, 3)
            out["has_low_mu"] = bool(finite_float(row["min_mu"], 1.0) < 0.99)
            cum_rows.append(out)
            start += length
    return lanes, pd.DataFrame(cum_rows)


def load_vehicle(relative_path: str, cache: dict[str, pd.DataFrame]) -> pd.DataFrame:
    if relative_path in cache:
        return cache[relative_path]
    path = RAW_ROOT / relative_path
    if not path.exists():
        cache[relative_path] = pd.DataFrame()
        return cache[relative_path]
    head = pd.read_csv(path, nrows=0, encoding="utf-8-sig")
    usecols = [c for c in VEHICLE_COLS if c in head.columns]
    df = pd.read_csv(path, usecols=usecols, encoding="utf-8-sig", low_memory=False)
    if "StorageTime" not in df.columns:
        cache[relative_path] = pd.DataFrame()
        return cache[relative_path]
    sec = to_seconds(df["StorageTime"])
    valid = np.isfinite(sec)
    df = df.loc[valid].copy()
    sec = sec[valid]
    if len(sec) == 0:
        cache[relative_path] = pd.DataFrame()
        return cache[relative_path]
    df["time_rel_s"] = sec - sec[0]
    for col in VEHICLE_COLS:
        if col != "StorageTime" and col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    cache[relative_path] = df
    return df


def value_at_peak(window: pd.DataFrame, col: str) -> tuple[float, float]:
    if col not in window.columns or window.empty:
        return float("nan"), float("nan")
    series = pd.to_numeric(window[col], errors="coerce")
    if series.notna().sum() == 0:
        return float("nan"), float("nan")
    idx = series.abs().idxmax()
    return finite_float(window.loc[idx, "time_rel_s"]), finite_float(window.loc[idx, col])


def value_at_peak_delta_from_start(window: pd.DataFrame, col: str) -> tuple[float, float]:
    if col not in window.columns or window.empty:
        return float("nan"), float("nan")
    series = pd.to_numeric(window[col], errors="coerce")
    valid = series.dropna()
    if valid.empty:
        return float("nan"), float("nan")
    baseline = float(valid.iloc[0])
    delta = series - baseline
    if delta.notna().sum() == 0:
        return float("nan"), float("nan")
    idx = delta.abs().idxmax()
    return finite_float(window.loc[idx, "time_rel_s"]), finite_float(delta.loc[idx])


def first_brake_time(window: pd.DataFrame) -> tuple[float, float]:
    if "zx|BrakePedal" not in window.columns or window.empty:
        return float("nan"), float("nan")
    brake = pd.to_numeric(window["zx|BrakePedal"], errors="coerce")
    if brake.notna().sum() == 0:
        return float("nan"), float("nan")
    max_brake = float(brake.max())
    threshold = 5.0 if max_brake > 1.5 else 0.05
    hit = window.loc[brake >= threshold]
    if hit.empty:
        return float("nan"), max_brake
    return finite_float(hit.iloc[0]["time_rel_s"]), max_brake


def mu_change_times(window: pd.DataFrame) -> list[dict[str, Any]]:
    if "zx1|mu" not in window.columns or window.empty:
        return []
    mu = pd.to_numeric(window["zx1|mu"], errors="coerce")
    valid = window.loc[mu.notna(), ["time_rel_s"]].copy()
    valid["mu"] = mu.loc[valid.index].to_numpy()
    if valid.empty:
        return []
    rows: list[dict[str, Any]] = []
    low = valid.loc[valid["mu"] < 0.99]
    if not low.empty:
        rows.append(
            {
                "candidate_anchor_type_cn": "原始车辆 mu 首次低附着",
                "candidate_time_rel_s": finite_float(low.iloc[0]["time_rel_s"]),
                "support_value": finite_float(low.iloc[0]["mu"]),
                "support_signal": "zx1|mu",
            }
        )
    last_mu = None
    for _, row in valid.iterrows():
        cur = finite_float(row["mu"])
        if last_mu is None:
            last_mu = cur
            continue
        if math.isfinite(cur) and math.isfinite(last_mu) and abs(cur - last_mu) >= 0.05:
            rows.append(
                {
                    "candidate_anchor_type_cn": "原始车辆 mu 跳变",
                    "candidate_time_rel_s": finite_float(row["time_rel_s"]),
                    "support_value": cur,
                    "support_signal": f"zx1|mu: {last_mu:.3f}->{cur:.3f}",
                }
            )
            last_mu = cur
    dedup: list[dict[str, Any]] = []
    seen: set[tuple[str, int]] = set()
    for row in rows:
        key = (row["candidate_anchor_type_cn"], int(round(row["candidate_time_rel_s"] * 10)))
        if key not in seen:
            seen.add(key)
            dedup.append(row)
    return dedup


def add_candidate(
    rows: list[dict[str, Any]],
    segment: pd.Series,
    module_rule: dict[str, str],
    anchor_type: str,
    candidate_time: float,
    source: str,
    support_signal: str = "",
    support_value: float = float("nan"),
    note: str = "",
) -> None:
    if not math.isfinite(candidate_time):
        return
    rows.append(
        {
            "subject": segment.get("subject", ""),
            "session_stamp": segment.get("session_stamp", ""),
            "vehicle_raw_relative_path": segment.get("vehicle_raw_relative_path", ""),
            "module_name": segment.get("module_name", ""),
            "instance_name": segment.get("instance_name", ""),
            "segment_index": segment.get("segment_index", ""),
            "segment_entry_time_rel_s": finite_float(segment.get("entry_time_rel_s")),
            "segment_exit_time_rel_s": finite_float(segment.get("exit_time_rel_s")),
            "segment_duration_s": finite_float(segment.get("duration_s")),
            "segment_mapping_reliability": segment.get("segment_mapping_reliability", ""),
            "design_category_cn": module_rule.get("design_category_cn", ""),
            "paper_evidence_cn": module_rule.get("paper_evidence_cn", ""),
            "primary_anchor_rule_cn": module_rule.get("primary_anchor_rule_cn", ""),
            "candidate_anchor_type_cn": anchor_type,
            "candidate_time_rel_s": round(candidate_time, 3),
            "candidate_offset_from_segment_entry_s": round(candidate_time - finite_float(segment.get("entry_time_rel_s")), 3),
            "candidate_source_cn": source,
            "support_signal": support_signal,
            "support_value": support_value,
            "anchor_status_cn": module_rule.get("anchor_status_cn", ""),
            "use_for_model_input": "否，当前只用于离线锚点重建/样本切片",
            "note_cn": note,
        }
    )


def segment_from_scene_trigger(trigger_row: pd.Series) -> pd.Series:
    entry = finite_float(trigger_row.get("segment_entry_time_rel_s"))
    exit_time = finite_float(trigger_row.get("segment_exit_time_rel_s"))
    duration = exit_time - entry if math.isfinite(entry) and math.isfinite(exit_time) else float("nan")
    return pd.Series(
        {
            "subject": trigger_row.get("subject", ""),
            "session_stamp": trigger_row.get("session_stamp", ""),
            "vehicle_raw_relative_path": trigger_row.get("vehicle_raw_relative_path", ""),
            "module_name": trigger_row.get("module_name", ""),
            "instance_name": trigger_row.get("instance_name", ""),
            "segment_index": "",
            "entry_time_rel_s": entry,
            "exit_time_rel_s": exit_time,
            "duration_s": duration,
            "segment_mapping_reliability": trigger_row.get("segment_mapping_reliability", ""),
        }
    )


def compact_display(value: Any) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return ""
    text = str(value).strip()
    try:
        number = float(text)
        if math.isfinite(number) and abs(number - round(number)) < 1e-9:
            return str(int(round(number)))
    except Exception:
        pass
    return text


def add_scene_trigger_candidates(rows: list[dict[str, Any]]) -> None:
    if not SCENE_TRIGGER_SESSION_TIMES.exists():
        return
    triggers = read_csv(SCENE_TRIGGER_SESSION_TIMES)
    if triggers.empty:
        return
    keep_modules = {"longstraight", "fix_road"}
    keep_triggers = {"ChangeLane", "Stop"}
    triggers = triggers[
        triggers["module_name"].astype(str).isin(keep_modules)
        & triggers["trigger_name"].astype(str).isin(keep_triggers)
    ].copy()
    if triggers.empty:
        return

    for _, trigger in triggers.iterrows():
        module = str(trigger.get("module_name", ""))
        rule = MODULE_RULES.get(module)
        if not rule:
            continue
        trigger_name = str(trigger.get("trigger_name", ""))
        if trigger_name == "Stop" and module != "longstraight":
            continue
        candidate_time = finite_float(trigger.get("estimated_trigger_time_rel_s"))
        segment = segment_from_scene_trigger(trigger)
        target_title = compact_display(trigger.get("target_title", ""))
        lane = compact_display(trigger.get("target_lane_id", "") or trigger.get("trigger_lane_id", ""))
        target_lane = compact_display(trigger.get("change_target_lane", ""))
        if trigger_name == "ChangeLane":
            anchor_type = "显式变道触发点"
            lane_text = f"{lane}->{target_lane}" if target_lane else lane
        else:
            anchor_type = "显式停车触发点"
            lane_text = lane
        support = f"{target_title} lane {lane_text}".strip()
        relative_s = finite_float(trigger.get("trigger_relative_s_in_module_m"))
        note = (
            f"来自 SILAB .aed 触发映射；trigger={trigger_name}，"
            f"target={target_title}，lane={lane_text}，"
            f"scene_trigger_uid={trigger.get('scene_trigger_uid', '')}。"
        )
        add_candidate(
            rows,
            segment,
            rule,
            anchor_type,
            candidate_time,
            "SILAB .aed 显式触发点",
            support,
            relative_s,
            note=note,
        )


def build_candidate_anchors(mu_segments: pd.DataFrame) -> pd.DataFrame:
    sessions = read_csv(SESSION_MODULE_TIMES)
    vehicle_cache: dict[str, pd.DataFrame] = {}
    rows: list[dict[str, Any]] = []
    mu_by_module: dict[str, pd.DataFrame] = {
        str(module): group.copy() for module, group in mu_segments.groupby("module_name")
    } if not mu_segments.empty else {}

    for _, segment in sessions.iterrows():
        module = str(segment.get("module_name", ""))
        rule = MODULE_RULES.get(module)
        if not rule:
            continue
        entry = finite_float(segment.get("entry_time_rel_s"))
        exit_time = finite_float(segment.get("exit_time_rel_s"))
        duration = finite_float(segment.get("duration_s"))
        if not (math.isfinite(entry) and math.isfinite(exit_time) and exit_time > entry):
            continue

        df = load_vehicle(str(segment.get("vehicle_raw_relative_path", "")), vehicle_cache)
        window = pd.DataFrame()
        if not df.empty:
            window = df[(df["time_rel_s"] >= entry) & (df["time_rel_s"] <= exit_time)].copy()

        if module == "longstraight":
            add_candidate(
                rows,
                segment,
                rule,
                "场景上下文入口",
                entry,
                "道路模块进入时间",
                note="当前不作为主事件锚点；用于记录连续驾驶/背景交通上下文。",
            )
            continue

        if module == "middle_section":
            add_candidate(rows, segment, rule, "连续超车段入口", entry, "道路连接段进入时间")
            add_candidate(rows, segment, rule, "连续超车段中点", entry + duration * 0.5, "道路连接段位置")
            t_lat, v_lat = value_at_peak_delta_from_start(window, "zx1|lateraldistance")
            add_candidate(rows, segment, rule, "横向偏移变化峰值", t_lat, "连续超车车身轨迹确认", "zx1|lateraldistance_delta", v_lat)
            t_ay, v_ay = value_at_peak(window, "zx|ay")
            add_candidate(rows, segment, rule, "横向加速度峰值", t_ay, "连续超车车身姿态确认", "zx|ay", v_ay)
            t_vyaw, v_vyaw = value_at_peak(window, "zx|vyaw")
            add_candidate(rows, segment, rule, "横摆角速度峰值", t_vyaw, "连续超车车身姿态确认", "zx|vyaw", v_vyaw)
            continue

        add_candidate(rows, segment, rule, "道路模块入口", entry, "道路设计位置")
        add_candidate(rows, segment, rule, "道路模块中点", entry + duration * 0.5, "道路设计位置")

        if module in {"curve1", "curve2", "curve3", "fix_road", "zd"}:
            t_roll, v_roll = value_at_peak(window, "zx|vroll")
            add_candidate(rows, segment, rule, "车身横滚速率峰值", t_roll, "车身姿态确认", "zx|vroll", v_roll)
            t_ay, v_ay = value_at_peak(window, "zx|ay")
            add_candidate(rows, segment, rule, "横向加速度峰值", t_ay, "车身姿态确认", "zx|ay", v_ay)
            t_vyaw, v_vyaw = value_at_peak(window, "zx|vyaw")
            add_candidate(rows, segment, rule, "横摆角速度峰值", t_vyaw, "车身姿态确认", "zx|vyaw", v_vyaw)

        if module == "differentmu_road":
            for item in mu_change_times(window):
                add_candidate(
                    rows,
                    segment,
                    rule,
                    item["candidate_anchor_type_cn"],
                    item["candidate_time_rel_s"],
                    "原始车辆路面附着字段",
                    item["support_signal"],
                    item["support_value"],
                )
            module_mu = mu_by_module.get(module, pd.DataFrame())
            if not module_mu.empty:
                total_len = float(module_mu["module_relative_s_end_m"].max())
                if total_len > 0:
                    last_min_mu = None
                    for _, mrow in module_mu.iterrows():
                        min_mu = finite_float(mrow.get("min_mu"))
                        rel_s = finite_float(mrow.get("module_relative_s_start_m"))
                        if rel_s <= 0:
                            last_min_mu = min_mu
                            continue
                        if last_min_mu is None or abs(min_mu - last_min_mu) >= 0.05:
                            t = entry + duration * (rel_s / total_len)
                            add_candidate(
                                rows,
                                segment,
                                rule,
                                "cfg 低附着段变化点",
                                t,
                                "道路配置推算",
                                "cfg lane mu",
                                min_mu,
                                note=f"按 cfg 相对距离 {rel_s:.1f}m / {total_len:.1f}m 线性换算到记录时间。",
                            )
                        last_min_mu = min_mu

        if module == "stop":
            t_brake, max_brake = first_brake_time(window)
            add_candidate(rows, segment, rule, "首次明显制动", t_brake, "车辆纵向响应确认", "zx|BrakePedal", max_brake)
            t_ax, v_ax = value_at_peak(window, "zx|ax")
            add_candidate(rows, segment, rule, "纵向加速度峰值", t_ax, "车辆纵向响应确认", "zx|ax", v_ax)

    add_scene_trigger_candidates(rows)
    return pd.DataFrame(rows)


def build_static_module_summary(candidates: pd.DataFrame, road_map: pd.DataFrame, mu_segments: pd.DataFrame) -> pd.DataFrame:
    candidate_counts = (
        candidates.groupby(["module_name", "candidate_anchor_type_cn"]).size().reset_index(name="count")
        if not candidates.empty
        else pd.DataFrame(columns=["module_name", "candidate_anchor_type_cn", "count"])
    )
    count_text = defaultdict(list)
    for _, row in candidate_counts.iterrows():
        count_text[str(row["module_name"])].append(f"{row['candidate_anchor_type_cn']}={row['count']}")

    rows: list[dict[str, Any]] = []
    for module, rule in MODULE_RULES.items():
        road = road_map[road_map["module_name"].astype(str) == module] if not road_map.empty else pd.DataFrame()
        mu = mu_segments[mu_segments["module_name"].astype(str) == module] if not mu_segments.empty else pd.DataFrame()
        rows.append(
            {
                "module_name": module,
                "design_category_cn": rule["design_category_cn"],
                "paper_evidence_cn": rule["paper_evidence_cn"],
                "primary_anchor_rule_cn": rule["primary_anchor_rule_cn"],
                "anchor_status_cn": rule["anchor_status_cn"],
                "road_instances": ";".join(road.get("instance_name", pd.Series(dtype=str)).astype(str).unique().tolist()),
                "road_s_range_m": ""
                if road.empty
                else f"{float(road['s_start_m'].min()):.1f}-{float(road['s_end_m'].max()):.1f}",
                "ego_direction_mu_values": ""
                if mu.empty
                else ";".join(str(v) for v in sorted(mu["min_mu"].dropna().unique())),
                "generated_candidate_counts": "；".join(count_text[module]),
                "next_action_cn": next_action_for_module(module),
            }
        )
    return pd.DataFrame(rows)


def next_action_for_module(module: str) -> str:
    if module == "longstraight":
        return "把 MAN TGL 25->26 显式变道和 Chrysler300 Stop 纳入候选锚点，再用车身横向/制动响应确认是否触发被试避让。"
    if module == "middle_section":
        return "按连续超车负荷事件段处理，优先检查连接段入口、横向偏移变化峰值、横摆/横向加速度峰值。"
    if module in {"curve1", "curve2", "curve3"}:
        return "检查局部横滚峰值/横向加速度峰值是否与道路曲率位置一致，再决定锚点。"
    if module == "differentmu_road":
        return "优先使用进入低 μ 区域和 μ 变化点；核对原始车辆 zx1|mu 与 cfg 是否一致。"
    if module == "fix_road":
        return "把 MAN TGL 25->26 和 BMW m340 26->25 两个显式变道触发纳入维修路段主候选，再结合车身姿态确认。"
    if module == "stop":
        return "确认是否存在被试方向前车急停；若没有显式触发，不能只靠背景 Stop。"
    if module == "zd":
        return "确认 zd 是否为匝道/汇入场景，并查汇入区设计位置。"
    return "继续确认该模块的具体实验设计语义。"


def write_report(
    design_extract_md: str,
    design_rows: pd.DataFrame,
    lane_rows: pd.DataFrame,
    mu_segments: pd.DataFrame,
    candidates: pd.DataFrame,
    module_summary: pd.DataFrame,
) -> None:
    (REPORT_DIR / "small_paper_scene_design_extract_v0_4.md").write_text(design_extract_md, encoding="utf-8")

    lines = [
        "# 被试方向设计点与候选锚点重建 v0.4",
        "",
        f"生成时间：{now_str()}",
        "",
        "## 这次做了什么",
        "",
        "本轮把用户提供的小论文作为实验设计文字依据，结合道路配置和已有车辆轨迹投影，生成“被试方向候选锚点”工作清单。",
        "",
        "重点不是训练模型，而是解决旧流程可能存在的样本语义错位：哪些点是真正被试方向事件，哪些只是背景交通或连续驾驶负荷。",
        "",
        "## 小论文提供的关键依据",
        "",
        "- 连续超车负荷：`middle_section` 重复 9 次，交通车约 23 m/s；用户补充确认道路连接段存在连续超车事件。",
        "- 大货车紧急变道：目标车侵入本车道，锚点依据为显式触发或方向盘角速度峰值。",
        "- 施工/维修路段：车道受限并伴随交通车干扰，预期变道避让和速度调整。",
        "- 前车急停：前车减速或停止，风险来源为 TTC 快速下降。",
        "- 低附着路段：`mu=0.8-0.2` 不同附着区域，锚点依据为进入低 μ 区域。",
        "- 弯道路段：`curve1/curve2/curve3`，风险来源为道路曲率变化，锚点依据为局部横滚峰值。",
        "- 匝道/汇入：多车道汇入与交通交互，锚点依据为汇入区或方向盘角速度峰值。",
        "",
        "## 当前生成的表格",
        "",
        f"- 小论文场景表：`{TABLE_DIR / 'small_paper_scene_design_tables_v0_4.csv'}`",
        f"- 配置车道/附着表：`{TABLE_DIR / 'cfg_lane_mu_geometry_v0_4.csv'}`",
        f"- 被试方向低附着段表：`{TABLE_DIR / 'cfg_ego_direction_mu_segments_v0_4.csv'}`",
        f"- 候选锚点清单：`{TABLE_DIR / 'ego_direction_design_anchor_candidates_v0_4.csv'}`",
        f"- 模块汇总表：`{TABLE_DIR / 'ego_direction_design_anchor_module_summary_v0_4.csv'}`",
        "",
        "## 数量概览",
        "",
        f"- 小论文场景/协议相关表格行数：{len(design_rows)}",
        f"- cfg 解析车道行数：{len(lane_rows)}",
        f"- 被试方向低附着/附着段行数：{len(mu_segments)}",
        f"- 候选锚点行数：{len(candidates)}",
        "",
        "按场景模块汇总：",
        "",
    ]
    if not module_summary.empty:
        lines.append("| 场景 | 设计类别 | 锚点状态 | 候选数量 | 下一步 |")
        lines.append("|---|---|---|---|---|")
        for _, row in module_summary.iterrows():
            lines.append(
                f"| `{row['module_name']}` | {row['design_category_cn']} | {row['anchor_status_cn']} | "
                f"{row['generated_candidate_counts']} | {row['next_action_cn']} |"
            )
    lines.extend(
        [
            "",
            "## 当前判断",
            "",
            "1. 现在可以确认：旧流程只用方向盘响应或泛化事件池来定锚点是不够稳的。",
            "2. 新锚点应分成两层：先有场景/道路/任务设计点，再用车身姿态确认是否产生响应。",
            "3. `longstraight` 不能再简单按背景交通处理：25/26 车流中包含 MAN TGL 显式变道和 Chrysler300 停车触发，应作为直道交通交互候选锚点。",
            "4. `fix_road` 也不是泛化维修场景：.aed 中已经能定位 MAN TGL 25->26 和 BMW m340 26->25 两个显式变道触发。",
            "5. 弯道、低附着、middle_section 连续超车、longstraight 显式交通触发、fix_road 显式变道都应进入下一轮锚点可视化审查。",
            "6. `stop`、`zd` 还需要继续查更具体的实验设计或用户说明，不能直接把背景交通触发当真值。",
            "",
            "## 下一步建议",
            "",
            "1. 先从本轮候选锚点中筛出弯道、低附着、middle_section 连续超车、longstraight 显式触发、fix_road 显式变道这些高优先级样本。",
            "2. 对 `stop`、`zd` 继续找显式实验设计位置，并核对是否位于被试实际行驶方向。",
            "3. 生成候选锚点可视化图，检查候选锚点前后车身姿态和方向盘响应是否符合物理意义。",
            "4. 在锚点视觉审查通过前，不建议继续用这些样本训练风格/生理模型。",
        ]
    )
    (REPORT_DIR / "ego_direction_design_anchor_rebuild_v0_4_cn.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    design_rows, design_md = extract_small_paper_design()
    lane_rows, mu_segments = parse_cfg_lane_geometry()
    road_map = read_csv(ROAD_POSITION_MAP) if ROAD_POSITION_MAP.exists() else pd.DataFrame()
    candidates = build_candidate_anchors(mu_segments)
    module_summary = build_static_module_summary(candidates, road_map, mu_segments)

    write_csv(design_rows, TABLE_DIR / "small_paper_scene_design_tables_v0_4.csv")
    write_csv(lane_rows, TABLE_DIR / "cfg_lane_mu_geometry_v0_4.csv")
    write_csv(mu_segments, TABLE_DIR / "cfg_ego_direction_mu_segments_v0_4.csv")
    write_csv(candidates, TABLE_DIR / "ego_direction_design_anchor_candidates_v0_4.csv")
    write_csv(module_summary, TABLE_DIR / "ego_direction_design_anchor_module_summary_v0_4.csv")
    write_report(design_md, design_rows, lane_rows, mu_segments, candidates, module_summary)

    print(
        {
            "design_rows": len(design_rows),
            "lane_rows": len(lane_rows),
            "mu_segments": len(mu_segments),
            "candidates": len(candidates),
            "modules": len(module_summary),
        }
    )


if __name__ == "__main__":
    main()
