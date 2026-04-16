from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ANSWER_PREFIXES = (
    "【答案】",
    "【解析】",
    "【导语】",
    "【详解】",
    "【点睛】",
    "【原文】",
)

ANALYSIS_NUMBER_PREFIXES = (
    "词汇积累",
    "句式拓展",
    "段落续写",
    "续写线索",
    "行为类",
    "情绪类",
)


def iter_paragraphs(doc: DocumentObject) -> list[Paragraph]:
    return list(doc.paragraphs)


def is_section_heading(text: str) -> bool:
    return bool(re.match(r"^第[一二三四五六七八九十]+部分", text)) or bool(
        re.match(r"^第[一二三四五六七八九十]+节", text)
    )


def is_article_label(text: str) -> bool:
    return text in {"A", "B", "C", "D"}


def is_instruction_start(text: str) -> bool:
    starters = (
        "听下面",
        "阅读下面",
        "做题时",
        "注意：",
        "注意:",
    )
    return text.startswith(starters)


def is_question_start(text: str) -> bool:
    match = re.match(r"^(\d+)\.\s*(.*)", text)
    if not match:
        return False
    rest = match.group(2).strip()
    if not rest:
        return False
    if any(rest.startswith(prefix) for prefix in ANALYSIS_NUMBER_PREFIXES):
        return False
    if rest.startswith("【"):
        return "音频" in rest or "此处可播放" in rest
    if re.search(r"\b\d+\.\s*", rest):
        return False
    if re.match(r"^[A-D]\.\s", rest):
        return True
    if rest.endswith(("?", "？")):
        return True
    if len(rest) >= 12 and re.search(r"[，。；：,.]", rest):
        return True
    return False


def is_resume_after_skip(text: str) -> bool:
    return (
        is_section_heading(text)
        or is_article_label(text)
        or is_instruction_start(text)
        or is_question_start(text)
    )


def extract_exam_paragraphs(src: DocumentObject) -> list[Paragraph]:
    kept: list[Paragraph] = []
    skipping = False

    for para in iter_paragraphs(src):
        text = para.text.strip()
        if not text:
            if not skipping and kept:
                kept.append(para)
            continue

        if text.startswith(ANSWER_PREFIXES):
            skipping = True
            continue

        if skipping:
            if is_resume_after_skip(text):
                skipping = False
            else:
                continue

        kept.append(para)

    while kept and not kept[-1].text.strip():
        kept.pop()
    return kept


def ensure_two_columns(section, column_space_cm: float) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols_el = cols[0]
    else:
        cols_el = OxmlElement("w:cols")
        sect_pr.append(cols_el)
    cols_el.set(qn("w:num"), "2")
    cols_el.set(qn("w:space"), str(int((Cm(column_space_cm) / 635))))


def set_run_fonts(run, size_pt: float) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")


def copy_runs(src_para: Paragraph, dst_para: Paragraph, size_pt: float) -> None:
    for src_run in src_para.runs:
        text = src_run.text
        if not text:
            continue
        dst_run = dst_para.add_run(text)
        dst_run.bold = src_run.bold
        dst_run.italic = src_run.italic
        dst_run.underline = src_run.underline
        set_run_fonts(dst_run, size_pt)

    if not src_para.runs:
        dst_run = dst_para.add_run(src_para.text)
        set_run_fonts(dst_run, size_pt)


def classify_paragraph(text: str, index: int) -> str:
    if index == 0:
        return "title_main"
    if index == 1:
        return "title_sub"
    if index == 2:
        return "title_meta"
    if is_section_heading(text):
        return "section"
    if is_article_label(text):
        return "article"
    if text.startswith(("第一节", "第二节")):
        return "subsection"
    if text.startswith(("做题时", "听下面", "阅读下面", "注意：", "注意:")):
        return "instruction"
    if text.startswith(("Dear ", "Yours", "Li Hua", "Paragraph 1", "Paragraph 2")):
        return "writing"
    return "body"


def apply_paragraph_style(para: Paragraph, kind: str, line_spacing: float, body_size: float) -> None:
    fmt = para.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = line_spacing

    if kind == "title_main":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_after = Pt(2)
        fmt.line_spacing = 1.0
    elif kind == "title_sub":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_after = Pt(2)
        fmt.line_spacing = 1.0
    elif kind == "title_meta":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_after = Pt(4)
        fmt.line_spacing = 1.0
    elif kind in {"section", "subsection"}:
        fmt.space_before = Pt(4)
        fmt.space_after = Pt(1)
        fmt.line_spacing = 1.0
        para.paragraph_format.keep_with_next = True
    elif kind == "article":
        fmt.space_before = Pt(4)
        fmt.space_after = Pt(1)
        fmt.line_spacing = 1.0
        para.paragraph_format.keep_with_next = True
    elif kind == "instruction":
        fmt.space_before = Pt(2)
        fmt.space_after = Pt(1)
    elif kind == "writing":
        fmt.space_before = Pt(1)
        fmt.space_after = Pt(1)
    else:
        if re.match(r"^\d+\.\s", para.text.strip()):
            fmt.space_before = Pt(1)
        if para.text.strip().startswith("__"):
            fmt.line_spacing = 1.0

    for run in para.runs:
        set_run_fonts(run, body_size)

    if kind == "title_main":
        for run in para.runs:
            set_run_fonts(run, body_size + 5)
            run.bold = True
    elif kind == "title_sub":
        for run in para.runs:
            set_run_fonts(run, body_size + 4)
            run.bold = True
    elif kind == "title_meta":
        for run in para.runs:
            set_run_fonts(run, body_size)
    elif kind in {"section", "subsection"}:
        for run in para.runs:
            set_run_fonts(run, body_size + 1)
            run.bold = True
    elif kind == "article":
        for run in para.runs:
            set_run_fonts(run, body_size + 1)
            run.bold = True


def clear_header_footer(doc: DocumentObject) -> None:
    for section in doc.sections:
        for part in (section.header, section.footer):
            for para in part.paragraphs:
                para.text = ""


def build_document(
    src_path: Path,
    out_path: Path,
    body_size: float,
    line_spacing: float,
    left_margin_cm: float,
    right_margin_cm: float,
    top_margin_cm: float,
    bottom_margin_cm: float,
    column_space_cm: float,
) -> int:
    src = Document(str(src_path))
    kept = extract_exam_paragraphs(src)

    dst = Document()
    clear_header_footer(dst)

    section = dst.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(42)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(left_margin_cm)
    section.right_margin = Cm(right_margin_cm)
    section.top_margin = Cm(top_margin_cm)
    section.bottom_margin = Cm(bottom_margin_cm)
    ensure_two_columns(section, column_space_cm)

    normal_style = dst.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(body_size)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.line_spacing = line_spacing
    normal_rpr = normal_style.element.rPr
    if normal_rpr is None:
        normal_rpr = OxmlElement("w:rPr")
        normal_style.element.append(normal_rpr)
    normal_fonts = normal_rpr.rFonts
    if normal_fonts is None:
        normal_fonts = OxmlElement("w:rFonts")
        normal_rpr.append(normal_fonts)
    normal_fonts.set(qn("w:ascii"), "Times New Roman")
    normal_fonts.set(qn("w:hAnsi"), "Times New Roman")
    normal_fonts.set(qn("w:cs"), "Times New Roman")
    normal_fonts.set(qn("w:eastAsia"), "宋体")

    dst._body.clear_content()

    for idx, src_para in enumerate(kept):
        text = src_para.text
        kind = classify_paragraph(text.strip(), idx)
        dst_para = dst.add_paragraph()
        copy_runs(src_para, dst_para, body_size)
        apply_paragraph_style(dst_para, kind, line_spacing, body_size)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    dst.save(str(out_path))
    return len(kept)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--body-size", type=float, default=10.5)
    parser.add_argument("--line-spacing", type=float, default=1.05)
    parser.add_argument("--left-margin-cm", type=float, default=1.3)
    parser.add_argument("--right-margin-cm", type=float, default=1.3)
    parser.add_argument("--top-margin-cm", type=float, default=1.1)
    parser.add_argument("--bottom-margin-cm", type=float, default=1.1)
    parser.add_argument("--column-space-cm", type=float, default=0.8)
    args = parser.parse_args()

    count = build_document(
        src_path=Path(args.source),
        out_path=Path(args.output),
        body_size=args.body_size,
        line_spacing=args.line_spacing,
        left_margin_cm=args.left_margin_cm,
        right_margin_cm=args.right_margin_cm,
        top_margin_cm=args.top_margin_cm,
        bottom_margin_cm=args.bottom_margin_cm,
        column_space_cm=args.column_space_cm,
    )
    print(f"kept_paragraphs={count}")
    print(f"output={args.output}")


if __name__ == "__main__":
    main()
