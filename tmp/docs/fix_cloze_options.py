from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table


QUESTION_RE = re.compile(
    r"^(?P<num>4[1-9]|5[0-5])\.\s*"
    r"A\.\s*(?P<A>.*?)\t"
    r"B\.\s*(?P<B>.*?)\t"
    r"C\.\s*(?P<C>.*?)\t"
    r"D\.\s*(?P<D>.*?)$"
)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))
    tc_w.set(qn("w:type"), "dxa")


def remove_table_borders(table: Table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")


def set_table_width(table: Table, width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))
    tbl_w.set(qn("w:type"), "dxa")


def set_run_font(run, size_pt: float) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")


def insert_table_after_paragraph(doc: Document, paragraph, texts: list[str]) -> Table:
    table = doc.add_table(rows=1, cols=5)
    tbl = table._tbl
    paragraph._p.addnext(tbl)

    widths = [1.2, 4.0, 4.0, 4.0, 4.0]
    row = table.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    for idx, (cell, text, width) in enumerate(zip(row.cells, texts, widths)):
        set_cell_width(cell, width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        if idx == 0:
            run.bold = False
        set_run_font(run, 11.7)

    table.autofit = False
    set_table_width(table, sum(widths))
    remove_table_borders(table)
    return table


def clear_paragraph(paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def main() -> None:
    src = Path(r"F:\育英学校高二英语周测三 柯瑜_试卷版.docx")
    tmp_src = Path(r"F:\data_set_process\data_process\tmp\docs\exam_current.docx")
    if src.exists():
        tmp_src.write_bytes(src.read_bytes())

    doc = Document(str(tmp_src))

    targets = []
    for para in doc.paragraphs:
        text = para.text.strip()
        m = QUESTION_RE.match(text)
        if m:
            targets.append((para, m))

    for para, m in reversed(targets):
        texts = [
            f"{m.group('num')}.",
            f"A. {m.group('A')}",
            f"B. {m.group('B')}",
            f"C. {m.group('C')}",
            f"D. {m.group('D')}",
        ]
        insert_table_after_paragraph(doc, para, texts)
        clear_paragraph(para)

    out = Path(r"F:\育英学校高二英语周测三 柯瑜_试卷版.docx")
    doc.save(str(out))
    print(f"updated_rows={len(targets)}")
    print(out)


if __name__ == "__main__":
    main()
